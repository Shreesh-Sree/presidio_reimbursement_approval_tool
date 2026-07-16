"""Safe, bounded text extraction for administrator-uploaded policy evidence."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path


class PolicyDocumentExtractionError(ValueError):
    """A policy file cannot supply text suitable for the evidence-only RAG index."""


def extract_policy_text(*, file_name: str, content: bytes) -> str:
    """Return a bounded text representation of a supported office document.

    The original file remains the system of record. This helper deliberately
    does not run arbitrary macros, invoke a shell converter, or OCR scanned
    PDFs; those need a sandboxed document-processing pipeline.
    """

    extension = Path(file_name).suffix.lower()
    if extension == ".pdf":
        try:
            from pypdf import PdfReader

            reader = PdfReader(BytesIO(content))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as exc:
            raise PolicyDocumentExtractionError("The PDF could not be read for policy indexing") from exc
    elif extension == ".docx":
        try:
            from docx import Document

            document = Document(BytesIO(content))
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            table_cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
            text = "\n".join((*paragraphs, *table_cells))
        except Exception as exc:
            raise PolicyDocumentExtractionError("The DOCX file could not be read for policy indexing") from exc
    else:
        raise PolicyDocumentExtractionError(
            "Automatic policy indexing supports text-based PDF and DOCX files. Paste approved text for this file type."
        )

    normalized = "\n".join(line.strip() for line in text.splitlines() if line.strip()).strip()
    if not normalized:
        raise PolicyDocumentExtractionError(
            "No selectable text was found. Upload a text-based PDF/DOCX or use the approved-text index option."
        )
    return normalized[:50_000]
