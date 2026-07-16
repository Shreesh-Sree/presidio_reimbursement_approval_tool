from io import BytesIO

import pytest
from docx import Document

from app.services.policy_document_text import PolicyDocumentExtractionError, extract_policy_text


def test_extracts_paragraphs_and_tables_from_docx_policy_evidence():
    document = Document()
    document.add_paragraph("Domestic travel is reimbursed at actual cost.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Meal cap"
    table.cell(0, 1).text = "INR 1,200"
    stream = BytesIO()
    document.save(stream)

    text = extract_policy_text(file_name="india-policy.docx", content=stream.getvalue())

    assert "Domestic travel" in text
    assert "INR 1,200" in text


def test_rejects_unsupported_legacy_documents_for_automatic_indexing():
    with pytest.raises(PolicyDocumentExtractionError, match="PDF and DOCX"):
        extract_policy_text(file_name="legacy-policy.xls", content=b"not-an-office-document")
